#!/usr/bin/env python3
"""
Word to Blog Converter
Converts Word documents (.docx) to Markdown blog posts for Astro/Skadoosh website
Extracts images and creates proper frontmatter
"""

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

try:
    from docx import Document
    from docx.shared import Inches
    import mammoth
    import requests
    from odf.opendocument import load
    from odf import text, draw
    from odf.element import Node
except ImportError:
    print("Required packages not installed. Run:")
    print("pip install python-docx mammoth requests pillow odfpy")
    sys.exit(1)


class WordToBlogConverter:
    def __init__(self, blog_root: Path):
        self.blog_root = Path(blog_root)
        self.content_dir = self.blog_root / "src" / "content" / "blog"
        self.public_dir = self.blog_root / "public" / "blog"
        
    def sanitize_filename(self, text: str) -> str:
        """Convert title to URL-friendly filename"""
        # Remove special characters and convert to lowercase
        text = re.sub(r'[^\w\s-]', '', text.strip())
        text = re.sub(r'[-\s]+', '-', text)
        return text.lower()
    
    def extract_images_mammoth(self, docx_path: Path, slug: str) -> tuple:
        """Extract images using mammoth (better HTML conversion)"""
        image_dir = self.public_dir / slug
        image_dir.mkdir(parents=True, exist_ok=True)
        
        image_counter = 1
        image_map = {}
        
        def convert_image(image):
            nonlocal image_counter
            
            # Generate filename
            extension = image.content_type.split('/')[-1]
            if extension == 'jpeg':
                extension = 'jpg'
            
            filename = f"image-{image_counter:02d}.{extension}"
            filepath = image_dir / filename
            
            # Save image
            with open(filepath, 'wb') as f:
                f.write(image.bytes)
            
            # Return relative path for markdown
            relative_path = f"/blog/{slug}/{filename}"
            image_map[image_counter] = relative_path
            image_counter += 1
            
            return {
                "src": relative_path,
                "alt": f"Image {image_counter - 1}"
            }
        
        # Convert document
        with open(docx_path, 'rb') as docx_file:
            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.img_element(convert_image)
            )
        
        return result.value, image_map, result.messages
    
    def html_to_markdown(self, html_content: str) -> str:
        """Convert HTML to Markdown (basic conversion)"""
        # Simple HTML to Markdown conversion
        # You might want to use a library like markdownify for better conversion
        
        # Headers
        html_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', html_content)
        html_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', html_content)
        html_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', html_content)
        html_content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'#### \1', html_content)
        
        # Bold and italic
        html_content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', html_content)
        html_content = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', html_content)
        html_content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', html_content)
        html_content = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', html_content)
        
        # Lists
        html_content = re.sub(r'<ul[^>]*>', '', html_content)
        html_content = re.sub(r'</ul>', '', html_content)
        html_content = re.sub(r'<ol[^>]*>', '', html_content)
        html_content = re.sub(r'</ol>', '', html_content)
        html_content = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1', html_content)
        
        # Images (keep as-is since we handled them in conversion)
        
        # Paragraphs
        html_content = re.sub(r'<p[^>]*>', '', html_content)
        html_content = re.sub(r'</p>', '\n\n', html_content)
        
        # Code blocks
        html_content = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', html_content)
        
        # Clean up extra whitespace
        html_content = re.sub(r'\n{3,}', '\n\n', html_content)
        html_content = html_content.strip()
        
        return html_content
    
    def extract_from_odt(self, odt_path: Path, slug: str) -> tuple:
        """Extract content and images from ODT file"""
        image_dir = self.public_dir / slug
        image_dir.mkdir(parents=True, exist_ok=True)
        
        doc = load(odt_path)
        content_parts = []
        image_counter = 1
        image_map = {}
        
        # Extract text content
        for element in doc.getElementsByType(text.P):
            paragraph_text = self.get_text_content(element)
            if paragraph_text.strip():
                # Simple formatting detection
                if any(style in str(element) for style in ['Heading', 'Title']):
                    # Try to determine heading level
                    if 'Heading 1' in str(element) or 'Title' in str(element):
                        content_parts.append(f"# {paragraph_text}")
                    elif 'Heading 2' in str(element):
                        content_parts.append(f"## {paragraph_text}")
                    elif 'Heading 3' in str(element):
                        content_parts.append(f"### {paragraph_text}")
                    else:
                        content_parts.append(f"## {paragraph_text}")
                else:
                    content_parts.append(paragraph_text)
                content_parts.append("")  # Add blank line
        
        # Extract images
        for element in doc.getElementsByType(draw.Image):
            try:
                href = element.getAttribute('href')
                if href:
                    # Extract image from ODT (it's a zip file)
                    import zipfile
                    with zipfile.ZipFile(odt_path, 'r') as zip_ref:
                        if href in zip_ref.namelist():
                            image_data = zip_ref.read(href)
                            
                            # Determine file extension
                            extension = href.split('.')[-1] if '.' in href else 'jpg'
                            filename = f"image-{image_counter:02d}.{extension}"
                            filepath = image_dir / filename
                            
                            with open(filepath, 'wb') as f:
                                f.write(image_data)
                            
                            relative_path = f"/blog/{slug}/{filename}"
                            image_map[image_counter] = relative_path
                            
                            # Insert image reference in content
                            content_parts.append(f"![Image {image_counter}]({relative_path})")
                            content_parts.append("")
                            
                            image_counter += 1
            except Exception as e:
                print(f"Warning: Could not extract image: {e}")
        
        return "\n".join(content_parts), image_map, []
    
    def get_text_content(self, element):
        """Recursively get text content from ODT element"""
        text_content = ""
        if hasattr(element, 'data'):
            text_content += str(element.data)
        for child in element.childNodes:
            if hasattr(child, 'data'):
                text_content += str(child.data)
            else:
                text_content += self.get_text_content(child)
        return text_content
    
    def extract_metadata_from_odt(self, odt_path: Path) -> dict:
        """Extract title and metadata from ODT document"""
        metadata = {
            'title': 'Untitled Blog Post',
            'description': '',
            'author': 'Suyash Pandey',
            'pubDate': datetime.now().strftime('%Y-%m-%d')
        }
        
        doc = load(odt_path)
        
        # Try to get title from first paragraph or heading
        for element in doc.getElementsByType(text.P):
            paragraph_text = self.get_text_content(element).strip()
            if paragraph_text:
                if len(paragraph_text) < 100:
                    metadata['title'] = paragraph_text
                    break
        
        # Generate description from first substantial paragraph
        for element in doc.getElementsByType(text.P):
            paragraph_text = self.get_text_content(element).strip()
            if paragraph_text and paragraph_text != metadata['title'] and len(paragraph_text) > 50:
                desc = paragraph_text[:150]
                if len(desc) < len(paragraph_text):
                    desc += "..."
                metadata['description'] = desc
                break
        
        return metadata
        """Extract title and metadata from Word document"""
        metadata = {
            'title': 'Untitled Blog Post',
            'description': '',
            'author': 'Suyash Pandey',
            'pubDate': datetime.now().strftime('%Y-%m-%d')
        }
        
        # Try to get title from first heading or first paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.style.name.startswith('Heading') or len(paragraph.text) < 100:
                    metadata['title'] = paragraph.text.strip()
                    break
                else:
                    # Use first sentence as title if no heading found
                    first_sentence = paragraph.text.split('.')[0].strip()
                    if len(first_sentence) < 80:
                        metadata['title'] = first_sentence
                    break
        
        # Generate description from first paragraph that's not the title
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() and paragraph.text.strip() != metadata['title']:
                desc = paragraph.text.strip()[:150]
                if len(desc) < len(paragraph.text.strip()):
                    desc += "..."
                metadata['description'] = desc
                break
        
        return metadata
    
    def create_frontmatter(self, metadata: dict, slug: str, has_images: bool) -> str:
        """Create Astro frontmatter"""
        hero_image = f"/blog/{slug}/image-01.jpg" if has_images else "/blog-placeholder-about.jpg"
        
        frontmatter = f"""---
title: "{metadata['title']}"
description: "{metadata['description']}"
pubDate: "{metadata['pubDate']}"
heroImage: "{hero_image}"
---"""
        
        return frontmatter
    
    def convert_docx_to_blog(self, doc_path: Path, output_slug: str = None) -> Path:
        """Main conversion function for both .docx and .odt files"""
        print(f"Converting {doc_path} to blog post...")
        
        file_extension = doc_path.suffix.lower()
        
        if file_extension == '.docx':
            return self.convert_docx_file(doc_path, output_slug)
        elif file_extension == '.odt':
            return self.convert_odt_file(doc_path, output_slug)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def convert_docx_file(self, docx_path: Path, output_slug: str = None) -> Path:
        """Convert DOCX file to blog post"""
        # Load document for metadata
        doc = Document(docx_path)
        metadata = self.extract_metadata(doc)
        
        # Generate slug
        if not output_slug:
            output_slug = self.sanitize_filename(metadata['title'])
        
        print(f"Generated slug: {output_slug}")
        print(f"Title: {metadata['title']}")
        
        # Extract content and images
        html_content, image_map, messages = self.extract_images_mammoth(docx_path, output_slug)
        
        # Show any warnings
        if messages:
            print("Conversion messages:")
            for message in messages:
                print(f"  {message}")
        
        # Convert to markdown
        markdown_content = self.html_to_markdown(html_content)
        
        return self.create_blog_file(metadata, output_slug, markdown_content, image_map)
    
    def convert_odt_file(self, odt_path: Path, output_slug: str = None) -> Path:
        """Convert ODT file to blog post"""
        # Extract metadata
        metadata = self.extract_metadata_from_odt(odt_path)
        
        # Generate slug
        if not output_slug:
            output_slug = self.sanitize_filename(metadata['title'])
        
        print(f"Generated slug: {output_slug}")
        print(f"Title: {metadata['title']}")
        
        # Extract content and images
        markdown_content, image_map, messages = self.extract_from_odt(odt_path, output_slug)
        
        return self.create_blog_file(metadata, output_slug, markdown_content, image_map)
    
    def create_blog_file(self, metadata: dict, output_slug: str, markdown_content: str, image_map: dict) -> Path:
        """Create frontmatter and write blog file"""
        # Create frontmatter
        frontmatter = self.create_frontmatter(metadata, output_slug, bool(image_map))
        
        # Combine frontmatter and content
        full_content = f"""{frontmatter}

{markdown_content}

---

**About the Author**

**Suyash Pandey** is the founder of Skadoosh B.V., specializing in Microsoft Dynamics 365 CE and Power Platform solutions. With over 10 years of experience in enterprise automation, Suyash helps organizations streamline their processes and improve efficiency through innovative Microsoft ecosystem implementations.

**Connect with Suyash:**
- [LinkedIn](https://www.linkedin.com/in/pandeysuyash/)
- [Email](mailto:suyash@ska-doosh.com)
- [Company Website](https://skadoosh-blogs.ska-doosh.workers.dev)
"""
        
        # Write markdown file
        output_file = self.content_dir / f"{output_slug}.md"
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(full_content)
        
        print(f"Created blog post: {output_file}")
        if image_map:
            print(f"Extracted {len(image_map)} images to: {self.public_dir / output_slug}")
        
        return output_file


def main():
    parser = argparse.ArgumentParser(description='Convert Word/ODT document to Astro blog post')
    parser.add_argument('doc_file', help='Path to document (.docx or .odt)')
    parser.add_argument('--slug', help='Custom URL slug (auto-generated if not provided)')
    parser.add_argument('--blog-root', default='.', help='Root directory of blog project')
    
    args = parser.parse_args()
    
    doc_path = Path(args.doc_file)
    if not doc_path.exists():
        print(f"Error: File {doc_path} not found")
        sys.exit(1)
    
    if not doc_path.suffix.lower() in ['.docx', '.odt']:
        print(f"Error: File must be a .docx or .odt file")
        sys.exit(1)
    
    blog_root = Path(args.blog_root).resolve()
    if not (blog_root / 'astro.config.mjs').exists():
        print(f"Error: {blog_root} doesn't appear to be an Astro blog project")
        sys.exit(1)
    
    converter = WordToBlogConverter(blog_root)
    output_file = converter.convert_docx_to_blog(doc_path, args.slug)
    
    print(f"\n✅ Conversion complete!")
    print(f"Blog post created: {output_file}")
    print(f"\nNext steps:")
    print(f"1. Review the generated markdown file")
    print(f"2. Run 'npm run build' to test")
    print(f"3. Run 'npm run preview' to preview")


if __name__ == "__main__":
    main()